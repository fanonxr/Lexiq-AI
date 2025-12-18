"""Document parsing service for various file types."""

import io
import re
from typing import Optional

from document_ingestion.config import get_settings
from document_ingestion.models.document import DocumentMetadata, ParsedDocument
from document_ingestion.utils.errors import ParsingError
from document_ingestion.utils.logging import get_logger

logger = get_logger("parser_service")
settings = get_settings()


class ParserService:
    """
    Service for parsing documents of various file types.

    Supports:
    - PDF (`.pdf`) - PyPDF2
    - DOCX (`.docx`) - python-docx
    - TXT (`.txt`) - Direct text extraction
    - MD (`.md`) - Markdown text
    """

    def __init__(self):
        """Initialize parser service."""
        self.allowed_types = settings.allowed_file_types

    async def parse_document(
        self, file_data: bytes, file_type: str, filename: Optional[str] = None
    ) -> ParsedDocument:
        """
        Parse document based on file type.

        Args:
            file_data: Raw file bytes
            file_type: File type/extension (pdf, docx, txt, md)
            filename: Optional filename for logging

        Returns:
            ParsedDocument with extracted text and metadata

        Raises:
            ParsingError: If parsing fails or file type is unsupported
        """
        file_type_lower = file_type.lower().strip(".")
        filename_str = filename or "unknown"

        # Validate file type
        if file_type_lower not in self.allowed_types:
            raise ParsingError(
                f"Unsupported file type: {file_type_lower}. "
                f"Allowed types: {', '.join(self.allowed_types)}",
                file_type=file_type_lower,
            )

        logger.info(f"Parsing document: type={file_type_lower}, filename={filename_str}")

        try:
            if file_type_lower == "pdf":
                return await self._parse_pdf(file_data, filename_str)
            elif file_type_lower == "docx":
                return await self._parse_docx(file_data, filename_str)
            elif file_type_lower in ["txt", "md"]:
                return await self._parse_text(file_data, file_type_lower, filename_str)
            else:
                raise ParsingError(
                    f"Parser not implemented for file type: {file_type_lower}",
                    file_type=file_type_lower,
                )
        except ParsingError:
            raise
        except Exception as e:
            logger.error(f"Unexpected error parsing document: {filename_str} - {e}", exc_info=True)
            raise ParsingError(
                f"Failed to parse document: {str(e)}",
                file_type=file_type_lower,
            ) from e

    async def _parse_pdf(self, file_data: bytes, filename: str) -> ParsedDocument:
        """
        Parse PDF file using PyPDF2.

        Args:
            file_data: PDF file bytes
            filename: Filename for logging

        Returns:
            ParsedDocument with extracted text and metadata

        Raises:
            ParsingError: If PDF parsing fails
        """
        try:
            import PyPDF2

            pdf_file = io.BytesIO(file_data)
            pdf_reader = PyPDF2.PdfReader(pdf_file)

            # Extract text from all pages
            text_parts = []
            for page_num, page in enumerate(pdf_reader.pages, start=1):
                try:
                    page_text = page.extract_text()
                    if page_text.strip():
                        text_parts.append(page_text)
                except Exception as page_error:
                    logger.warning(
                        f"Failed to extract text from page {page_num} in {filename}: {page_error}"
                    )
                    # Continue with other pages
                    continue

            if not text_parts:
                raise ParsingError(
                    "No text could be extracted from PDF. The file may be image-based or corrupted.",
                    file_type="pdf",
                )

            full_text = "\n\n".join(text_parts)

            # Extract metadata
            metadata = pdf_reader.metadata or {}
            page_count = len(pdf_reader.pages)

            # Calculate word count (approximate)
            word_count = len(re.findall(r"\b\w+\b", full_text))
            character_count = len(full_text)

            document_metadata = DocumentMetadata(
                page_count=page_count,
                word_count=word_count,
                character_count=character_count,
                title=metadata.get("/Title"),
                author=metadata.get("/Author"),
                created_at=str(metadata.get("/CreationDate")) if metadata.get("/CreationDate") else None,
                modified_at=str(metadata.get("/ModDate")) if metadata.get("/ModDate") else None,
            )

            logger.info(
                f"Successfully parsed PDF: {filename}, "
                f"pages={page_count}, words={word_count}, chars={character_count}"
            )

            return ParsedDocument(
                text=full_text,
                metadata=document_metadata,
                file_type="pdf",
            )

        except PyPDF2.errors.PdfReadError as e:
            raise ParsingError(
                f"PDF file is corrupted or invalid: {str(e)}",
                file_type="pdf",
            ) from e
        except Exception as e:
            raise ParsingError(
                f"Failed to parse PDF: {str(e)}",
                file_type="pdf",
            ) from e

    async def _parse_docx(self, file_data: bytes, filename: str) -> ParsedDocument:
        """
        Parse DOCX file using python-docx.

        Args:
            file_data: DOCX file bytes
            filename: Filename for logging

        Returns:
            ParsedDocument with extracted text and metadata

        Raises:
            ParsingError: If DOCX parsing fails
        """
        try:
            from docx import Document

            docx_file = io.BytesIO(file_data)
            doc = Document(docx_file)

            # Extract text from all paragraphs
            text_parts = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)

            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        text_parts.append(row_text)

            if not text_parts:
                raise ParsingError(
                    "No text could be extracted from DOCX file. The file may be empty or corrupted.",
                    file_type="docx",
                )

            full_text = "\n\n".join(text_parts)

            # Extract metadata
            core_props = doc.core_properties
            word_count = len(re.findall(r"\b\w+\b", full_text))
            character_count = len(full_text)

            document_metadata = DocumentMetadata(
                word_count=word_count,
                character_count=character_count,
                title=core_props.title,
                author=core_props.author,
                created_at=core_props.created.isoformat() if core_props.created else None,
                modified_at=core_props.modified.isoformat() if core_props.modified else None,
            )

            logger.info(
                f"Successfully parsed DOCX: {filename}, "
                f"words={word_count}, chars={character_count}"
            )

            return ParsedDocument(
                text=full_text,
                metadata=document_metadata,
                file_type="docx",
            )

        except Exception as e:
            raise ParsingError(
                f"Failed to parse DOCX: {str(e)}",
                file_type="docx",
            ) from e

    async def _parse_text(
        self, file_data: bytes, file_type: str, filename: str
    ) -> ParsedDocument:
        """
        Parse plain text or markdown file.

        Args:
            file_data: Text file bytes
            file_type: File type (txt or md)
            filename: Filename for logging

        Returns:
            ParsedDocument with extracted text and metadata

        Raises:
            ParsingError: If text parsing fails
        """
        try:
            # Try UTF-8 first, then fallback to other encodings
            encodings = ["utf-8", "latin-1", "cp1252", "iso-8859-1"]

            text = None
            used_encoding = None

            for encoding in encodings:
                try:
                    text = file_data.decode(encoding)
                    used_encoding = encoding
                    break
                except UnicodeDecodeError:
                    continue

            if text is None:
                raise ParsingError(
                    "Failed to decode text file. Unsupported encoding.",
                    file_type=file_type,
                )

            # Remove BOM if present
            if text.startswith("\ufeff"):
                text = text[1:]

            if not text.strip():
                raise ParsingError(
                    "Text file is empty.",
                    file_type=file_type,
                )

            # Calculate word count
            word_count = len(re.findall(r"\b\w+\b", text))
            character_count = len(text)

            document_metadata = DocumentMetadata(
                word_count=word_count,
                character_count=character_count,
                encoding=used_encoding,
            )

            logger.info(
                f"Successfully parsed {file_type.upper()}: {filename}, "
                f"words={word_count}, chars={character_count}, encoding={used_encoding}"
            )

            return ParsedDocument(
                text=text,
                metadata=document_metadata,
                file_type=file_type,
            )

        except ParsingError:
            raise
        except Exception as e:
            raise ParsingError(
                f"Failed to parse text file: {str(e)}",
                file_type=file_type,
            ) from e

